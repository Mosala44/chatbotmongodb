from django.shortcuts import render, redirect
from .models import camiones_collection, datos_collection
from django.http import JsonResponse, HttpResponse
from .db_connection import db
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime
from docx.shared import Inches
import matplotlib.pyplot as plt
from datetime import datetime

motor_info = {
        "2CAM3080": {"MT1": "W09060763", "MT2": "W09040922"},
        "2CAM3082": {"MT1": "W09070380", "MT2": ""},
        "2CAM3083": {"MT1": "W08040386", "MT2": "W11030045"},
        "2CAM3085": {"MT1": "W11070959", "MT2": "WX15030035"},
        "2CAM3086": {"MT1": "W11050868", "MT2": "W11090094"},
        "2CAM3087": {"MT1": "W09040921", "MT2": ""},
        "2CAM3090": {"MT1": "W06010096", "MT2": "W13010323"},
        "2CAM3091": {"MT1": "W11060945", "MT2": "W08070121"},
        "2CAM3092": {"MT1": "W12070264", "MT2": "Sin placa"}
    }

thresholds = {
    "viscocidad": {"monitoreo": 90, "accion": 90},  # Solo se evalúa acción: >90
    
    "fe": {"monitoreo": 50, "accion": 100},
    "cu": {"monitoreo": 8, "accion": 15},
    "pb": {"monitoreo": 7, "accion": 10},
    "al": {"monitoreo": 8, "accion": 15},
    "sn": {"monitoreo": 3, "accion": 5},
    "cr": {"monitoreo": 2, "accion": 4},
    "ni": {"monitoreo": 2, "accion": 4},
    "si": {"monitoreo": 30, "accion": 80},
    "na": {"monitoreo": None, "accion": 2},
    "cfe": {"monitoreo": 50, "accion": 100}
}


def set_cell_shading(cell, fill_color):
    """
    Establece el color de fondo de una celda.
    fill_color: cadena hexadecimal, ej. "FFFF00" para amarillo.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    tcPr = cell._tc.get_or_add_tcPr()
    # Eliminar sombreado previo, si existe
    for child in tcPr.findall(qn('w:shd')):
        tcPr.remove(child)
    tcPr.append(shading_elm)

def set_row_shading(row_cells, fill_color):
    """
    Aplica el sombreado solo a la fila completa.
    """
    for cell in row_cells:
        set_cell_shading(cell, fill_color)


def index(request):
    camiones = datos_collection.find({})
    data = {camiones: "camiones"}
    return render (request, "index.html", data)
# Create your views here.
def create_analisis(data: dict):
    return datos_collection.insert_one(data).inserted_id
def lista_camiones(request):
    """Obtiene todos los camiones y los muestra en la lista."""
    camiones = list(camiones_collection.find({}, {"_id": 0, "numero": 1}))  # Solo traemos el campo 'numero'
    return render(request, "lista.html", {"camiones": camiones})

def create_camion(request):
    """Crea un nuevo camión y redirige a la lista."""
    if request.method == "POST":
        numero_camion = request.POST.get("numero_camion")

        if numero_camion:
            # Verificar si el camión ya existe
            if camiones_collection.find_one({"numero": numero_camion}):
                return redirect("lista_camiones")  # Redirigir a la lista de camiones si ya existe

            # Insertar nuevo camión
            camiones_collection.insert_one({"numero": numero_camion})

            return redirect("lista_camiones")  # Redirigir a la lista de camiones tras agregarlo

    return redirect("lista_camiones")   # Si el formulario no es válido, redirigir igual


from django.http import JsonResponse

PREGUNTAS = [
    "número de camión: ",
    "número de muestra: ",
    "motor_1: ",
    "motor_2: ",
    "tipo de pauta: ",
    "horas componente: ",
    "fecha análisis (dia-mes-año): ",
    "cambio de aceite: ",
    "lubricant: ",
    "horómetro: ",
    "fecha de la muestra (dia-mes-año): ",
    "viscocidad: ",
    "agua %: ",
    "cfe: ",
    "fe: ",
    "cu: ",
    "pb: ",
    "al: ",
    "sn: ",
    "ag: ",
    "cr: ",
    "ni: ",
    "mo: ",
    "ti: ",
    "si: ",
    "na: ",
    "k: ",
    "b: ",
    "v: ",
    "mg: ",
    "ca: ",
    "p: ",
    "zn: ",
    "ba: ",
    "cd: ",
    "li: ",
    "mn: ",
    "sb: "
]

def selecciona_camion(request):
    """
    Vista para seleccionar un camión antes de iniciar el análisis.
    Se guarda el camión seleccionado en la sesión y se inicializa el diccionario de datos.
    """
    camiones_collection = db["camion"]
    if request.method == "POST":
        camion_id = request.POST.get("camion")
        if camion_id:
            # Buscamos en la colección "camion" usando el campo "numero"
            camion = camiones_collection.find_one({"numero": camion_id})
            if camion:
                # Guardamos en la sesión el número del camión
                request.session["selected_camion"] = camion["numero"]
                # Inicializamos el diccionario de análisis con el camión ya seleccionado
                request.session["analisis_data"] = {"Número de camión:": camion["numero"]}
                print("✅ Camión almacenado en sesión:", camion["numero"])
                return redirect('cht')
            else:
                # Si el camión no existe, se muestra un error
                camiones = list(camiones_collection.find({}))
                return render(request, 'selec_cam.html', {
                    'error': 'El camión seleccionado no existe.',
                    'camiones': camiones
                })
    # Método GET: se muestra la lista de camiones para elegir
    camiones = list(camiones_collection.find({}))
    return render(request, 'selec_cam.html', {'camiones': camiones})
def reiniciar_chat(request):
    """
    Vista para reiniciar la conversación del chat.
    Conserva el camión seleccionado y borra las respuestas previas,
    de modo que la siguiente pregunta sea "número de muestra:".
    """
    # Recuperar el camión seleccionado desde la sesión
    camion_numero = request.session.get("selected_camion")
    if not camion_numero:
        return JsonResponse({"error": "No hay camión seleccionado"}, status=400)
    
    # Reiniciar la conversación conservando el camión seleccionado
    request.session["analisis_data"] = {"Número de camión:": camion_numero}
    
    # Retornar la respuesta inicial: la siguiente pregunta (número de muestra)
    # Asumiendo que en tu lista de preguntas, "número de camión:" es la primera,
    # la siguiente (segunda) es "número de muestra:"
    return JsonResponse({"message": "Chat reiniciado. Por favor, ingrese el número de muestra:"})



def Chatbot(request): 
    if request.method != "POST":
        return JsonResponse({"error": "Método no permitido"}, status=405)

    # Obtener la entrada del usuario y los datos actuales de la sesión
    user_input = request.POST.get("user_input", "").strip()
    session_data = request.session.get("analisis_data", {})

    # Recuperar el número de camión desde la sesión (ya se guardó en selecciona_camion)
    camion_numero = request.session.get("selected_camion")
    print(f"🔍 Debug - Número de camión en sesión: {camion_numero}")
    print(f"📦 Datos actuales en sesión: {session_data}")

    # Si el camión ya está en la sesión, no es necesario volver a preguntar
    # Por lo tanto, usamos la lista de preguntas a partir de la segunda
    if "Número de camión:" in session_data:
        preguntas = PREGUNTAS[1:]
    else:
        preguntas = PREGUNTAS

    # Calcular el avance en la conversación (restamos 1 si ya se tiene el camión)
    avance = len(session_data) - (1 if "Número de camión:" in session_data else 0)

    # Si aún no se han respondido todas las preguntas, guardamos la respuesta actual
    if avance < len(preguntas):
        session_data[preguntas[avance]] = user_input
        request.session["analisis_data"] = session_data

    # Si se han respondido todas las preguntas, se procesa y se guarda el análisis
    if avance + 1 == len(preguntas):
        try:
            analisis_data = {
                "camion": {"numero": camion_numero},  # Se usa el camión seleccionado
                "numero_muestra": session_data.get("número de muestra: ", ""),
                "motor_1": session_data.get("motor_1: ", "false").lower() in ["true", "si", "1"],
                "motor_2": session_data.get("motor_2: ", "false").lower() in ["true", "si", "1"],
                "tipo_pauta": session_data.get("tipo de pauta: ", ""),
                "horas_componentes": session_data.get("horas componente: ", ""),
                "fecha_analisis": session_data.get("fecha análisis (dia-mes-año): ", ""),
                "cambio_aceite": session_data.get("cambio de aceite: ", ""),
                "lubricant": session_data.get("lubricant: ", ""),
                "horometro": session_data.get("horómetro: ", ""),
                "fecha_muestra": session_data.get("fecha de la muestra (dia-mes-año): ", ""),
                "viscocidad": session_data.get("viscosidad: ", ""),
                "agua": session_data.get("agua %: ", ""),
                "cfe": session_data.get("cfe: ", ""),
                "fe": session_data.get("fe: ", ""),
                "cu": session_data.get("cu: ", ""),
                "pb": session_data.get("pb: ", ""),
                "al": session_data.get("al: ", ""),
                "sn": session_data.get("sn: ", ""),
                "ag": session_data.get("ag: ", ""),
                "cr": session_data.get("cr: ", ""),
                "ni": session_data.get("ni: ", ""),
                "mo": session_data.get("mo: ", ""),
                "ti": session_data.get("ti: ", ""),
                "si": session_data.get("si: ", ""),
                "na": session_data.get("na: ", ""),
                "k": session_data.get("k: ", ""),
                "b": session_data.get("b: ", ""),
                "v": session_data.get("v: ", ""),
                "mg": session_data.get("mg: ", ""),
                "ca": session_data.get("ca: ", ""),
                "p": session_data.get("p: ", ""),
                "zn": session_data.get("zn: ", ""),
                "ba": session_data.get("ba: ", ""),
                "cd": session_data.get("cd: ", ""),
                "li": session_data.get("li: ", ""),
                "mn": session_data.get("mn: ", ""),
                "sb": session_data.get("b: ", "")
            }

            print("🚚 Datos a guardar en MongoDB:", analisis_data)

            def evaluar_condicion(elemento, valor):
                 rangos = {
                     "viscocidad": (90, 90),  # Mismo valor para monitoreo y acción si no hay rango definido
                     "cfe": (50, 100),
                     "fe": (50, 100),
                     "cu": (8, 15),
                     "pb": (7, 10),
                     "al": (8, 15),
                     'sn': (3, 5),
                     'cr': (2, 4),
                     'ni': (2, 4),
                     "si": (30, 80),
                     "na": (2, 2),  # Solo tiene NORMAL y ACCIÓN REQUERIDA
                 }
                 if elemento in rangos:
                     monitoreo, accion = rangos[elemento]
                     if valor > accion:
                         return "ACCIÓN REQUERIDA"
                     elif valor > monitoreo:
                         return "MONITOREO"
                 return "NORMAL"
             
             # Evaluar todos los elementos convirtiendo a float
            estados = {elem: evaluar_condicion(elem, float(analisis_data.get(elem, 0) or 0)) for elem in ["viscocidad", "fe", "cu", "pb", "al", "sn","cr", "ni", "si", "na"]}
             
             # Determinar la condición global del camión
            prioridad = {"NORMAL": 0, "MONITOREO": 1, "ACCIÓN REQUERIDA": 2}
            estado_final = "NORMAL"
            for estado in estados.values():
                 if prioridad[estado] > prioridad[estado_final]:
                     estado_final = estado
             
            
            # Insertar el documento en la colección "datos"
            datos_collection.insert_one(analisis_data)
            
            # Reiniciamos la sesión para la conversación, conservando el camión seleccionado
            request.session["analisis_data"] = {"Número de camión:": camion_numero}
            
            mensaje_final = f"La condición del análisis de este camión es: {estado_final}"
            
            return JsonResponse({"message": f"Datos guardados exitosamente 🎉\n\n{mensaje_final}"})
        except Exception as e:
            return JsonResponse({"error": f"Error al guardar: {str(e)}"}, status=500)
    else:
        # Si aún quedan preguntas, se devuelve la siguiente pregunta
        return JsonResponse({"message": preguntas[avance + 1]})


def to_float(val):
    try:
        return float(val)
    except:
        return 0.0

def evaluar_condicion(elemento, valor):
    rangos = {
        "viscocidad": (90, 90),  # Mismo valor para monitoreo y acción si no hay rango definido
        "cfe": (50, 100),
        "fe": (50, 100),
        "cu": (8, 15),
        "pb": (7, 10),
        "al": (8, 15),
        'sn': (3, 5),
        'cr': (2, 4),
        'ni': (2, 4),
        "si": (30, 80),
        "na": (2, 2),  # Solo tiene NORMAL y ACCIÓN REQUERIDA
    }
    if elemento in rangos:
        try:
            # Convertir valor a float (maneja decimales y enteros)
            valor_numerico = float(valor)
        except (ValueError, TypeError):
            return "NORMAL"  # Si no es convertible, se considera normal
        
        monitoreo, accion = rangos[elemento]
        
        if valor_numerico > accion:
            return "ACCIÓN REQUERIDA"
        elif valor_numerico > monitoreo:
            return "MONITOREO"
    return "NORMAL"

# Evaluar todos los elementos usando to_float para la conversión
def generar_recomendaciones(analisis_m1, analisis_m2, estados):
    recomendaciones = []

    # Verificar si el camión tiene análisis en ambos motores
    if not analisis_m1 or not analisis_m2:
        recomendaciones.append("Sin registro en planilla de análisis.")

    # Verificar actualización de horómetro
    horas_m1 = analisis_m1.get("horas_componentes", "no hay registros") if analisis_m1 else "no hay registros"
    horas_m2 = analisis_m2.get("horas_componentes", "no hay registros") if analisis_m2 else "no hay registros"

    actualizar_horometro = []
    if horas_m1 in ["0", "no hay registros"]:
        actualizar_horometro.append("Motor 1")
    if horas_m2 in ["0", "no hay registros"]:
        actualizar_horometro.append("Motor 2")

    if actualizar_horometro:
        recomendaciones.append(f"Actualizar horómetro ({', '.join(actualizar_horometro)})")

    # Verificar Silicio (Si) en MONITOREO o ACCIÓN REQUERIDA
    if estados.get("si") in ["MONITOREO", "ACCIÓN REQUERIDA"]:
        recomendaciones.append(
            "Revisar motivo de incremento de sílices, tales como revisar sellos de tapas de inspección de piñón y llenado de aceite, "
            "además de realizar chequeo de mangueras de respiraderos de MT."
        )

    # Verificar si hay al menos 2 elementos en MONITOREO
    elementos_monitoreo = [elem for elem, estado in estados.items() if estado == "MONITOREO"]
    if len(elementos_monitoreo) >= 2:
        recomendaciones.extend([
            "Mantener el camión en monitoreo.",
            "Realizar chequeo y registro fotográfico del tapón magnético del cárter.",
            "Realizar chequeo y registro fotográfico del piñón solar y dientes de los planetarios.",
            "Revisar el estado del filtro de respiradero.",
            "Informar supervisión en caso de material escamoso o daños en piñones.",
            "Realizar una diálisis al aceite."
        ])

    # Verificar si hay al menos 2 elementos en ACCIÓN REQUERIDA
    elementos_accion = [elem for elem, estado in estados.items() if estado == "ACCIÓN REQUERIDA"]
    if len(elementos_accion) >= 2:
        recomendaciones.extend([
            "Realizar acciones correctivas inmediatas.",
            "Chequeo y registro fotográfico del tapón magnético y componentes internos.",
            "Drenado, flushing, limpieza exhaustiva y cambio del aceite.",
            "Si hay daños en piñón solar, realizar su cambio y torqueo correspondiente.",
            "Medir backlash y end play de los planetarios, informar anomalías."
        ])

    # Si todas las condiciones fueron "NORMAL", no se generan recomendaciones
    if all(estado == "NORMAL" for estado in estados.values()):
        return ["No se generan recomendaciones."]

    return recomendaciones


def generar_informe(request):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from django.http import HttpResponse
    from datetime import datetime
    from io import BytesIO
    # Asegúrate de tener importados los siguientes:
    # from docx.oxml import OxmlElement
    # from docx.oxml.ns import qn
    # import matplotlib.pyplot as plt

    # Obtener todos los camiones de la colección "camion"
    camiones = list(db["camion"].find({}))
    if not camiones:
        return HttpResponse("No hay camiones registrados.", status=400)
    
    # Eliminar duplicados basados en el campo "numero"
    unique_camiones = {camion["numero"]: camion for camion in camiones}
    camiones = list(unique_camiones.values())
    
    doc = Document()
    
    # Encabezado principal
    header = doc.sections[0].header
    
    # Agregar la tabla al encabezado (2 filas x 3 columnas)
    table_enc = header.add_table(rows=2, cols=3, width=Inches(6))
    table_enc.style = 'Table Grid'
    
    # Ruta de la imagen
    image_path = "C:/Users/Andres Villarroel/chatbotai2/static/images/image.png"
    
    # Insertar la imagen en la celda (0,0)
    cell_imagen = table_enc.cell(0, 0)
    paragraph = cell_imagen.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(1))
    
    # Agregar el texto en la primera fila
    table_enc.cell(0, 1).text = "Centro de Reparación de Componentes Antofagasta"
    table_enc.cell(0, 2).text = "ANÁLISIS SEMANAL DE ACEITE"
    
    # Segunda fila: Información adicional
    fecha_generacion = datetime.now().strftime("%d-%m-%Y")  # Formato DD-MM-YYYY
    table_enc.cell(1, 0).text = f"Fecha: {fecha_generacion}"
    
    cell_titulo = table_enc.cell(1, 1)
    p_titulo = cell_titulo.paragraphs[0]
    run_titulo = p_titulo.add_run("Reporte de análisis Muestra de Aceite")
    run_titulo.bold = True
    
    # Campo de paginación dinámico
    cell_paginas = table_enc.cell(1, 2)
    paginacion_paragraph = cell_paginas.paragraphs[0]
    field_page = OxmlElement('w:fldSimple')
    field_page.set(qn('w:instr'), 'PAGE')
    run_page = OxmlElement('w:r')
    text_page = OxmlElement('w:t')
    text_page.text = "X"  # Se actualizará con el número real de página
    run_page.append(text_page)
    field_page.append(run_page)
    paginacion_paragraph._element.append(field_page)
    paginacion_paragraph.add_run(" of ")
    field_numpages = OxmlElement('w:fldSimple')
    field_numpages.set(qn('w:instr'), 'NUMPAGES')
    run_numpages = OxmlElement('w:r')
    text_numpages = OxmlElement('w:t')
    text_numpages.text = "Y"  # Se actualizará con el total de páginas
    run_numpages.append(text_numpages)
    field_numpages.append(run_numpages)
    paginacion_paragraph._element.append(field_numpages)
    
    doc.add_paragraph("\n")
    
    # Agregar tabla de encabezado con Cliente y Fecha
    tabla_encabezado = doc.add_table(rows=1, cols=2)
    tabla_encabezado.style = 'Table Grid'
    
    celda_cliente = tabla_encabezado.cell(0, 0)
    p_cliente = celda_cliente.paragraphs[0]
    run_cliente = p_cliente.add_run("Cliente: ")
    run_cliente.bold = True
    run_cliente.font.size = Pt(12)
    p_cliente.add_run("__________________________")
    
    celda_fecha = tabla_encabezado.cell(0, 1)
    p_fecha = celda_fecha.paragraphs[0]
    run_fecha = p_fecha.add_run("Fecha: ")
    run_fecha.bold = True
    run_fecha.font.size = Pt(12)
    p_fecha.add_run("__________________________")
    
    # Alinear la tabla
    for row in tabla_encabezado.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    doc.add_paragraph("\n")
    
    # Tabla resumen principal
    table_datos = doc.add_table(rows=1, cols=8)
    table_datos.style = 'Table Grid'
    hdr_cells = table_datos.rows[0].cells
    encabezados = ['Camión', 'Número de muestra', 'Fecha de Análisis', 'MT1 Horas', 'MT2 Horas', 'Condición', 'Observación', 'Recomendaciones']
    for i, txt in enumerate(encabezados):
        hdr_cells[i].text = txt
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.bold = True
    
    datos_collection = db["datos"]
    critical_camiones = []
    elementos = ["viscocidad", "cfe", "fe", "cu", "pb", "al", "sn", "cr", "ni", "si", "na"]
    # Prioridad para evaluar condiciones
    prioridad = {"NORMAL": 0, "MONITOREO": 1, "ACCIÓN REQUERIDA": 2}
    
    # Procesar cada camión y agregar filas a la tabla resumen
    for camion in camiones:
        numero = camion.get("numero")
        analisis_m1 = datos_collection.find_one({"camion.numero": numero, "motor_1": True}, sort=[("numero_muestra", -1)])
        analisis_m2 = datos_collection.find_one({"camion.numero": numero, "motor_2": True}, sort=[("numero_muestra", -1)])
        horometro_m1 = str(analisis_m1.get("horometro", "N/A")) if analisis_m1 else "N/A"
        horometro_m2 = str(analisis_m2.get("horometro", "N/A")) if analisis_m2 else "N/A"
        
        if not analisis_m1 and not analisis_m2:
            continue
        
        estados_m1 = {}
        estados_m2 = {}
        motor_origen_param = {}
        
        for elem in elementos:
            key = elem.lower()
            val1 = to_float(analisis_m1.get(key, 0)) if analisis_m1 else 0
            val2 = to_float(analisis_m2.get(key, 0)) if analisis_m2 else 0
            
            estados_m1[elem] = evaluar_condicion(elem, val1)
            estados_m2[elem] = evaluar_condicion(elem, val2)
            # Determina de cuál motor proviene el dato mayor
            motor_origen_param[elem] = 1 if val1 >= val2 else 2

        estado_final_m1 = "NORMAL"
        estado_final_m2 = "NORMAL"
        for estado in estados_m1.values():
            if prioridad[estado.upper()] > prioridad[estado_final_m1.upper()]:
                estado_final_m1 = estado
        for estado in estados_m2.values():
            if prioridad[estado.upper()] > prioridad[estado_final_m2.upper()]:
                estado_final_m2 = estado
        
        fecha_m1 = analisis_m1.get("fecha_analisis", "") if analisis_m1 else ""
        horas_componentes_m1 = str(analisis_m1.get("horas_componentes", "N/A")) if analisis_m1 else "N/A"
        nro_muestra_m1 = analisis_m1.get("numero_muestra", "") if analisis_m1 else ""
        
        fecha_m2 = analisis_m2.get("fecha_analisis", "") if analisis_m2 else ""
        horas_componentes_m2 = str(analisis_m2.get("horas_componentes", "N/A")) if analisis_m2 else "N/A"
        nro_muestra_m2 = analisis_m2.get("numero_muestra", "") if analisis_m2 else ""
        
        fecha_comb = f"{fecha_m1} / {fecha_m2}" if fecha_m1 and fecha_m2 else (fecha_m1 or fecha_m2)
        nro_muestra_comb = f"{nro_muestra_m1} / {nro_muestra_m2}" if nro_muestra_m1 and nro_muestra_m2 else (nro_muestra_m1 or nro_muestra_m2)
        
        estado_global = estado_final_m1 if prioridad[estado_final_m1.upper()] > prioridad[estado_final_m2.upper()] else estado_final_m2
        
        obs_items = []
        for elem in elementos:
            if estados_m1[elem] != "NORMAL":
                obs_items.append(f"Motor 1: {elem.upper()} ({analisis_m1.get(elem.lower(), 0)})")
            if estados_m2[elem] != "NORMAL":
                obs_items.append(f"Motor 2: {elem.upper()} ({analisis_m2.get(elem.lower(), 0)})")
        observacion_global = ", ".join(obs_items)
        estados = {elem: max(estados_m1[elem], estados_m2[elem], key=lambda x: prioridad[x]) for elem in elementos}
        recomendaciones = generar_recomendaciones(analisis_m1, analisis_m2, estados)
        
        row_cells = table_datos.add_row().cells
        row_cells[0].text = str(numero)
        row_cells[1].text = nro_muestra_comb
        row_cells[2].text = fecha_comb
        row_cells[3].text = horas_componentes_m1
        row_cells[4].text = horas_componentes_m2
        row_cells[5].text = estado_global
        row_cells[6].text = observacion_global
        row_cells[7].text = "recomendaciones en graficos"
        
        if estado_global.upper() == "MONITOREO":
            set_row_shading(row_cells, "FFFF00")
        elif estado_global.upper() == "ACCIÓN REQUERIDA":
            set_row_shading(row_cells, "FF0000")
        
        # Registrar camiones críticos (no NORMAL)
        if estado_global.upper() != "NORMAL":
            critical_camiones.append({
                "numero": numero,
                # Se guarda además la información de los análisis para cada motor
                "analisis_m1": analisis_m1,
                "analisis_m2": analisis_m2,
                "horas_componentes_m1": horas_componentes_m1,
                "horas_componentes_m2": horas_componentes_m2,
                "estados_m1": estados_m1,  # <- Debe ser el diccionario completo
                "estados_m2": estados_m2,  # <- Debe ser el diccionario completo
                "estado_final_m1": estado_final_m1,
                "estado_final_m2": estado_final_m2,
                "motor_origen_param": motor_origen_param
            })
    
    # =========================================================================
    # Sección: Listas de camiones críticos por motor (sin gráficos ni tablas de gráficos)
    #
    # Se desea generar dos listas (una para Motor 1 y otra para Motor 2)
    # que incluyan, para cada camión crítico, los elementos: cfe, fe, cu, si, cr, ni
    # en estado MONITOREO o ACCIÓN REQUERIDA.
    # Cada entrada debe contener:
    #   - Número del camión
    #   - Número de muestra (último registro; se consultarán los últimos 5)
    #   - Motor (Motor 1 o Motor 2)
    #   - Fechas de análisis y número de muestra de los últimos 5 registros
    #   - Elemento crítico
    # =========================================================================
    
    # Listas para cada motor
    lista_motor1 = []
    lista_motor2 = []
    criticos = ["cfe", "fe", "cu", "si", "cr", "ni"]
    
    for camion_data in critical_camiones:
        numero = camion_data["numero"]
        # Por cada parámetro crítico evaluado para cada motor
        for elem in criticos:
            # Verificar para Motor 1
            estado_m1 = camion_data["estados_m1"].get(elem, "NORMAL").upper()
            if estado_m1 in ["MONITOREO", "ACCIÓN REQUERIDA"]:
                # Consultar últimos 5 registros para Motor 1
                registros = list(datos_collection.find({"camion.numero": numero, "motor_1": True}).sort([("fecha_analisis", 1), ("numero_muestra", 1)]))
                ultimos = registros[:5] if len(registros) >= 5 else registros
                registros_info = [{
                    "numero_muestra": rec.get("numero_muestra", ""),
                    "fecha_analisis": rec.get("fecha_analisis", "")[:10]
                } for rec in ultimos]
                lista_motor1.append({
                    "camion": numero,
                    "motor": "Motor 1",
                    "elemento_critico": elem.upper(),
                    "ultimos_registros": registros_info
                })
            # Verificar para Motor 2
            estado_m2 = camion_data["estados_m2"].get(elem, "NORMAL").upper()
            if estado_m2 in ["MONITOREO", "ACCIÓN REQUERIDA"]:
                registros = list(datos_collection.find({"camion.numero": numero, "motor_2": True}).sort([("fecha_analisis", 1), ("numero_muestra", 1)]))
                ultimos = registros[:5] if len(registros) >= 5 else registros
                registros_info = [{
                    "numero_muestra": rec.get("numero_muestra", ""),
                    "fecha_analisis": rec.get("fecha_analisis", "")[:10]
                } for rec in ultimos]
                lista_motor2.append({
                    "camion": numero,
                    "motor": "Motor 2",
                    "elemento_critico": elem.upper(),
                    "ultimos_registros": registros_info
                })
    
    # Agregar las listas al documento
    # Agregar las listas al documento
    doc.add_page_break()
    

    elementos_criticos = ["cfe", "fe", "cu", "si", "cr", "ni"]
   
    
    for camion_data in critical_camiones:
        numero_camion = camion_data["numero"]  # <- Usar el número del camión actual
        tiene_elementos_criticos = False
        
        # Chequear Motor 1
        if camion_data["analisis_m1"]:
            for elem in elementos_criticos:
                if camion_data["estados_m1"].get(elem, "NORMAL") != "NORMAL":
                    tiene_elementos_criticos = True
                    break
        
        # Chequear Motor 2
        if not tiene_elementos_criticos and camion_data["analisis_m2"]:
            for elem in elementos_criticos:
                if camion_data["estados_m2"].get(elem, "NORMAL") != "NORMAL":
                    tiene_elementos_criticos = True
                    break
        
        # Solo generar contenido si hay elementos críticos
        if not tiene_elementos_criticos:
            continue  # Saltar a siguiente camión
        
        
        # Crear nueva página para cada camión
        doc.add_page_break()
        
        # Título del camión (usar numero_camion)
      

        # Obtener información de motores (usar numero_camion)
        numero_camion_str = str(numero_camion)  # Asegurar que sea string
        motores = motor_info.get(numero_camion_str, {"MT1": "N/A", "MT2": "N/A"})
        
        # Asignar valores (manejar casos vacíos o "Sin placa")
        mt1_serial = motores["MT1"] if motores["MT1"] else "N/A"
        mt2_serial = motores["MT2"] if motores["MT2"] else "N/A"

        motor_info_data = {
            "MT1": mt1_serial,
            "MT2": mt2_serial
        }

        # Crear tabla de 4 filas x 2 columnas
        titulo = doc.add_paragraph()
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titulo.add_run(f"INFORME DETALLADO - CAMIÓN {numero_camion}")
        run.bold = True
        run.font.size = Pt(14)
        # En la sección donde se genera la tabla
        tabla_info = doc.add_table(rows=4, cols=2)
        tabla_info.style = 'Table Grid'
        
        # Motor 1
        celda = tabla_info.cell(0, 0)
        celda.text = "Motor 1"
        celda = tabla_info.cell(1, 0)
        celda.text = f"N° Serie: {motor_info_data.get('MT1', 'N/A')}"  # Corregido
        celda = tabla_info.cell(2, 0)
        celda.text = f"Estado: {camion_data.get('estado_final_m1', 'N/A')}"  # Clave corregida
        celda = tabla_info.cell(3, 0)
        horas_componentes_m1 = (camion_data['horas_componentes_m1']  # Cambiado a horometro_m1
                       if camion_data['horas_componentes_m1'] not in ['N/A', 'Sin registro', '0', 'no hay registros']
                       else 'Actualizar horómetro')
        celda.text = f"Horómetro: {horas_componentes_m1} hrs"
        
        # Motor 2
        celda = tabla_info.cell(0, 1)
        celda.text = "Motor 2"
        celda = tabla_info.cell(1, 1)
        celda.text = f"N° Serie: {motor_info_data.get('MT2', 'N/A')}"  # Corregido
        celda = tabla_info.cell(2, 1)
        celda.text = f"Estado: {camion_data.get('estado_final_m2', 'N/A')}"  # Clave corregida
        celda = tabla_info.cell(3, 1)
        horas_componentes_m2 = (camion_data['horas_componentes_m2']  # Cambiado a horometro_m2
                       if camion_data['horas_componentes_m2'] not in ['N/A', 'Sin registro', '0', 'no hay registros']
                       else 'Actualizar horómetro')
        celda.text = f"Horómetro: {horas_componentes_m2} hrs"
        
        # Aplicar colores a los estados
        for col in [0, 1]:
            estado_cell = tabla_info.cell(2, col)
            estado = estado_cell.text.split(": ")[1]
            
            if estado == "ACCIÓN REQUERIDA":
                set_cell_shading(estado_cell, "FF0000")  # Rojo
            elif estado == "MONITOREO":
                set_cell_shading(estado_cell, "FFFF00")  # Amarillo
        

        # Estilizar la tabla
        for row in tabla_info.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)

        doc.add_paragraph("\n")
        
        # Procesar Motor 1
        if camion_data["analisis_m1"]:
            
            
            for elemento in elementos_criticos:
                if camion_data["estados_m1"].get(elemento, "NORMAL") != "NORMAL":
                    plt.figure(figsize=(8, 5))
                    plt.title(f"Camion {numero_camion} - Motor 1\nEvolución de {elemento.upper()}")
                    plt.xlabel("Fecha de análisis (Número de muestra)")
                    plt.ylabel("Valor (ppm)")
                    plt.grid(True)

                    # Obtener registros históricos
                    registros = list(datos_collection.find({
                        "camion.numero": numero_camion,
                        "motor_1": True,
                        elemento: {"$exists": True}
                    }))

                    # Eliminar duplicados y procesar fechas
                    registros_unicos = {}
                    for r in registros:
                        try:
                            fecha_str = r["fecha_analisis"][:10]
                            fecha = datetime.strptime(fecha_str, "%d-%m-%Y")
                            nro_muestra = r.get("numero_muestra", "S/N")
                            clave_unica = f"{fecha_str}-{nro_muestra}"
                            
                            if clave_unica not in registros_unicos:
                                registros_unicos[clave_unica] = {
                                    "fecha": fecha,
                                    "muestra": nro_muestra,
                                    "valor": to_float(r.get(elemento, 0))
                                }
                        except Exception as e:
                            print(f"Error procesando registro: {e}")
                            continue

                    # Ordenar registros por fecha
                    registros_ordenados = sorted(
                        registros_unicos.values(),
                        key=lambda x: x["fecha"]
                    )

                    # Tomar últimos 5 registros (los más recientes)
                    ultimos_registros = registros_ordenados[-5:] if len(registros_ordenados) >= 5 else registros_ordenados

                    # Preparar datos para el gráfico
                    etiquetas = [f"{r['fecha'].strftime('%d-%m-%Y')}\n({r['muestra']})" for r in ultimos_registros]
                    valores = [r['valor'] for r in ultimos_registros]

                    # Trazar gráfico
                    plt.plot(etiquetas, valores, marker='o', linestyle='-', color='blue', markersize=8)
                    
                    # Añadir líneas de referencia
                    thresholds_elemento = thresholds.get(elemento, {})
                    if thresholds_elemento.get("monitoreo"):
                        plt.axhline(y=thresholds_elemento["monitoreo"], color='orange', linestyle='--', label='Límite Monitoreo')
                    if thresholds_elemento.get("accion"):
                        plt.axhline(y=thresholds_elemento["accion"], color='red', linestyle='-', label='Límite Acción')

                    plt.legend()
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()
                    
                    # Guardar y añadir al documento
                    buf = BytesIO()
                    plt.savefig(buf, format='png', bbox_inches='tight')
                    buf.seek(0)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(buf, width=Inches(6))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    plt.close()

        # Procesar Motor 2 (misma lógica que Motor 1)
        if camion_data["analisis_m2"]:
            
            for elemento in elementos_criticos:
                if camion_data["estados_m2"].get(elemento, "NORMAL") != "NORMAL":
                    plt.figure(figsize=(8, 5))
                    plt.title(f"Camion {numero_camion} - Motor 2\nEvolución de {elemento.upper()}")
                    plt.xlabel("Fecha de análisis (Número de muestra)")
                    plt.ylabel("Valor (ppm)")
                    plt.grid(True)

                    # Obtener registros históricos
                    registros = list(datos_collection.find({
                        "camion.numero": numero_camion,
                        "motor_2": True,
                        elemento: {"$exists": True}
                    }))

                    # Eliminar duplicados y procesar fechas
                    registros_unicos = {}
                    for r in registros:
                        try:
                            fecha_str = r["fecha_analisis"][:10]
                            fecha = datetime.strptime(fecha_str, "%d-%m-%Y")
                            nro_muestra = r.get("numero_muestra", "S/N")
                            clave_unica = f"{fecha_str}-{nro_muestra}"
                            
                            if clave_unica not in registros_unicos:
                                registros_unicos[clave_unica] = {
                                    "fecha": fecha,
                                    "muestra": nro_muestra,
                                    "valor": to_float(r.get(elemento, 0))
                                }
                        except Exception as e:
                            print(f"Error procesando registro: {e}")
                            continue

                    # Ordenar registros por fecha
                    registros_ordenados = sorted(
                        registros_unicos.values(),
                        key=lambda x: x["fecha"]
                    )

                    # Tomar últimos 5 registros (los más recientes)
                    ultimos_registros = registros_ordenados[-5:] if len(registros_ordenados) >= 5 else registros_ordenados

                    # Preparar datos para el gráfico
                    etiquetas = [f"{r['fecha'].strftime('%d-%m-%Y')}\n({r['muestra']})" for r in ultimos_registros]
                    valores = [r['valor'] for r in ultimos_registros]

                    # Trazar gráfico
                    plt.plot(etiquetas, valores, marker='o', linestyle='-', color='blue', markersize=8)
                    
                    # Añadir líneas de referencia
                    thresholds_elemento = thresholds.get(elemento, {})
                    if thresholds_elemento.get("monitoreo"):
                        plt.axhline(y=thresholds_elemento["monitoreo"], color='orange', linestyle='--', label='Límite Monitoreo')
                    if thresholds_elemento.get("accion"):
                        plt.axhline(y=thresholds_elemento["accion"], color='red', linestyle='-', label='Límite Acción')

                    plt.legend()
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()
                    
                    # Guardar y añadir al documento
                    buf = BytesIO()
                    plt.savefig(buf, format='png', bbox_inches='tight')
                    buf.seek(0)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(buf, width=Inches(6))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    plt.close()


        doc.add_paragraph("\n")
    if elemento in ["si", "fe", "cfe"]:
       estado_elemento = estados_motor.get(elemento, "NORMAL").upper()  # estados_motor = estados_m1 o estados_m2
    
    if estado_elemento in ["MONITOREO", "ACCIÓN REQUERIDA"]:
        # Crear tabla para recomendaciones
        tabla_rec = doc.add_table(rows=1, cols=1)
        tabla_rec.style = "Table Grid"
        celda_titulo = tabla_rec.cell(0, 0)
        celda_titulo.text = f"Recomendaciones para {elemento.upper()} (Motor {motor_num})"
        
        # Añadir fila con viñetas
        row_rec = tabla_rec.add_row().cells[0]
        
        # Generar recomendaciones según elemento
        recomendaciones = []
        if elemento == "si":
            recomendaciones = [
                "Micro filtrado/cambio de aceite",
                "Chequear tapas de llenado e inspección",
                "Verificar mangueras respiradero",
                "Revisar filtro de respiradero MT",
                "Registro fotográfico piñón y planetarios",
                "Medición Backlash y EndPlay"
            ]
        elif elemento == "fe":
            recomendaciones = [
                "Micro filtrado/cambio de aceite",
                "Registro fotográfico componentes",
                "Inspección piñón solar"
            ]
            # Lógica adicional para horas
            if horas_componentes_motor >= 8000:  # horas_componentes_motor obtenidas del análisis
                recomendaciones.append("Metrología/cambio piñón (horas altas)")
        elif elemento == "cfe":
            recomendaciones = [
                "Micro filtrado/cambio de aceite",
                "Inspección tapón magnético",
                "Programar medición Backlash"
            ]
        
        # Añadir viñetas
        for rec in recomendaciones:
            p = row_rec.add_paragraph(rec, style="List Bullet")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Estilizar tabla
        for cell in tabla_rec.columns[0].cells:
            cell.width = Inches(6)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph() 
    
        # Ajuste de fuente en la tabla resumen (a partir de la fila 1)
    for row in table_datos.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="informe_analisis_aceite.docx"'
    doc.save(response)
    return response
