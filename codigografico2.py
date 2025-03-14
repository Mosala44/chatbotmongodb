def generar_informe(request):
    from docx import Document
    from docx.shared import Inches, Pt
    from django.http import HttpResponse
    from datetime import datetime
    from io import BytesIO

    # Obtener todos los camiones de la colección "camion"
    camiones = list(db["camion"].find({}))
    if not camiones:
        return HttpResponse("No hay camiones registrados.", status=400)
    
    # Eliminar duplicados basados en el campo "numero"
    unique_camiones = {camion["numero"]: camion for camion in camiones}
    camiones = list(unique_camiones.values())
    
    doc = Document()
    
    # Encabezado principal
    header = doc.sections[0].header
    
    # Agregar la tabla al encabezado (3 filas x 3 columnas)
    table_enc = header.add_table(rows=2, cols=3, width=Inches(6))
    table_enc.style = 'Table Grid'
    
    # Ruta de la imagen
    image_path = "C:/Users/Andres Villarroel/chatbotai2/static/images/image.png"
    
    # Insertar la imagen en la celda (0,0)
    cell_imagen = table_enc.cell(0, 0)
    paragraph = cell_imagen.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(1))
    
    # Agregar el texto en la primera fila
    table_enc.cell(0, 1).text = "Centro de Reparación de Componentes Antofagasta"
    table_enc.cell(0, 2).text = "ANÁLISIS SEMANAL DE ACEITE"
    
    # Segunda fila: Nueva información agregada
    # 1️⃣ Fecha de generación del informe
    fecha_generacion = datetime.now().strftime("%d-%m-%Y")  # Formato DD-MM-YYYY
    table_enc.cell(1, 0).text = f"Fecha: {fecha_generacion}"
    
    # 2️⃣ Título "Reporte de análisis Muestra de Aceite" en negrita
    cell_titulo = table_enc.cell(1, 1)
    p_titulo = cell_titulo.paragraphs[0]
    run_titulo = p_titulo.add_run("Reporte de análisis Muestra de Aceite")
    run_titulo.bold = True  # Texto en negrita
    
    # 3️⃣ Índice de páginas dinámico (esto se actualizará después)
    cell_paginas = table_enc.cell(1, 2)
    paginacion_paragraph = cell_paginas.paragraphs[0]
    
    # Campo para mostrar la página actual (X)
    field_page = OxmlElement('w:fldSimple')
    field_page.set(qn('w:instr'), 'PAGE')
    run_page = OxmlElement('w:r')
    text_page = OxmlElement('w:t')
    text_page.text = "X"  # Se reemplazará con el número real de página
    run_page.append(text_page)
    field_page.append(run_page)
    paginacion_paragraph._element.append(field_page)
    
    # Agregar " of "
    paginacion_paragraph.add_run(" of ")
    
    # Campo para mostrar el total de páginas (Y)
    field_numpages = OxmlElement('w:fldSimple')
    field_numpages.set(qn('w:instr'), 'NUMPAGES')
    run_numpages = OxmlElement('w:r')
    text_numpages = OxmlElement('w:t')
    text_numpages.text = "Y"  # Se reemplazará con el total de páginas
    run_numpages.append(text_numpages)
    field_numpages.append(run_numpages)
    paginacion_paragraph._element.append(field_numpages)
    

    doc.add_paragraph("\n")

    
# Combinar la tercera fila en una sola celda
   
    
    # Agregar tabla de encabezado con Cliente y Fecha
    tabla_encabezado = doc.add_table(rows=1, cols=2)
    tabla_encabezado.style = 'Table Grid'
    
    # Configurar primera celda (Cliente)
    celda_cliente = tabla_encabezado.cell(0, 0)
    p_cliente = celda_cliente.paragraphs[0]
    run_cliente = p_cliente.add_run("Cliente: ")
    run_cliente.bold = True
    run_cliente.font.size = Pt(12)
    p_cliente.add_run("__________________________")  # Espacio para escribir
    
    # Configurar segunda celda (Fecha)
    celda_fecha = tabla_encabezado.cell(0, 1)
    p_fecha = celda_fecha.paragraphs[0]
    run_fecha = p_fecha.add_run("Fecha: ")
    run_fecha.bold = True
    run_fecha.font.size = Pt(12)
    p_fecha.add_run("__________________________")  # Espacio para escribir
    
    # Alinear tabla al centro
    for row in tabla_encabezado.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Agregar espacio después del encabezado
    doc.add_paragraph("\n")


    # Tabla resumen
    table_datos = doc.add_table(rows=1, cols=8)
    table_datos.style = 'Table Grid'
    hdr_cells = table_datos.rows[0].cells
    encabezados = ['Camión', 'Número de muestra', 'Fecha de Análisis', 'MT1 Horas', 'MT2 Horas', 'Condición', 'Observación', 'Recomendaciones']
    for i, txt in enumerate(encabezados):
        hdr_cells[i].text = txt
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.bold = True
    datos_collection = db["datos"]
    critical_camiones = []
    elementos = ["viscocidad", "cfe", "fe", "cu", "pb", "al", "sn", "cr", "ni", "si", "na"]
    prioridad = {"NORMAL": 0, "MONITOREO": 1, "ACCIÓN REQUERIDA": 2}

    for camion in camiones:
        numero = camion.get("numero")
        analisis_m1 = datos_collection.find_one({"camion.numero": numero, "motor_1": True}, sort=[("numero_muestra", -1)])
        analisis_m2 = datos_collection.find_one({"camion.numero": numero, "motor_2": True}, sort=[("numero_muestra", -1)])
        
        if not analisis_m1 and not analisis_m2:
            continue

        # Procesar cada motor por separado
        estados_m1 = {}
        estados_m2 = {}
        motor_origen_param = {}
        
        for elem in elementos:
            key = elem.lower()
            val1 = to_float(analisis_m1.get(key, 0)) if analisis_m1 else 0
            val2 = to_float(analisis_m2.get(key, 0)) if analisis_m2 else 0
            
            estados_m1[elem] = evaluar_condicion(elem, val1)
            estados_m2[elem] = evaluar_condicion(elem, val2)
            motor_origen_param[elem] = 1 if val1 >= val2 else 2

        # Determinar estado final por motor
        estado_final_m1 = "NORMAL"
        estado_final_m2 = "NORMAL"
        for estado in estados_m1.values():
            if prioridad[estado.upper()] > prioridad[estado_final_m1.upper()]:
                estado_final_m1 = estado
        for estado in estados_m2.values():
            if prioridad[estado.upper()] > prioridad[estado_final_m2.upper()]:
                estado_final_m2 = estado

        # Datos comunes
        fecha_m1 = analisis_m1.get("fecha_analisis", "") if analisis_m1 else ""
        horas_componentes_m1 = str(analisis_m1.get("horas_componentes", "N/A")) if analisis_m1 else "N/A"
        nro_muestra_m1 = analisis_m1.get("numero_muestra", "") if analisis_m1 else ""
        
        fecha_m2 = analisis_m2.get("fecha_analisis", "") if analisis_m2 else ""
        horas_componentes_m2 = str(analisis_m2.get("horas_componentes", "N/A")) if analisis_m2 else "N/A"
        nro_muestra_m2 = analisis_m2.get("numero_muestra", "") if analisis_m2 else ""
        
        fecha_comb = f"{fecha_m1} / {fecha_m2}" if fecha_m1 and fecha_m2 else (fecha_m1 or fecha_m2)
        nro_muestra_comb = f"{nro_muestra_m1} / {nro_muestra_m2}" if nro_muestra_m1 and nro_muestra_m2 else (nro_muestra_m1 or nro_muestra_m2)
        
        # Determinar estado global
        estado_global = estado_final_m1 if prioridad[estado_final_m1.upper()] > prioridad[estado_final_m2.upper()] else estado_final_m2
        
        # Observaciones
        obs_items = []
        for elem in elementos:
            if estados_m1[elem] != "NORMAL":
                obs_items.append(f"Motor 1: {elem.upper()} ({analisis_m1.get(elem.lower(), 0)})")
            if estados_m2[elem] != "NORMAL":
                obs_items.append(f"Motor 2: {elem.upper()} ({analisis_m2.get(elem.lower(), 0)})")
        
        observacion_global = ", ".join(obs_items)
        estados = {elem: max(estados_m1[elem], estados_m2[elem], key=lambda x: prioridad[x]) for elem in elementos}
        recomendaciones = generar_recomendaciones(analisis_m1, analisis_m2, estados)
        
        # Agregar fila a la tabla
        row_cells = table_datos.add_row().cells
        row_cells[0].text = str(numero)
        row_cells[1].text = nro_muestra_comb
        row_cells[2].text = fecha_comb
        row_cells[3].text = horas_componentes_m1
        row_cells[4].text = horas_componentes_m2
        row_cells[5].text = estado_global
        row_cells[6].text = observacion_global
        row_cells[7].text = "recomendaciones en graficos"

        if estado_global.upper() == "MONITOREO":
            set_row_shading(row_cells, "FFFF00")
        elif estado_global.upper() == "ACCIÓN REQUERIDA":
            set_row_shading(row_cells, "FF0000")

        if estado_global.upper() != "NORMAL":
            critical_camiones.append({
                "numero": numero,
                "estado_m1": estado_final_m1,
                "estado_m2": estado_final_m2,
                "horas_componentes_m1": horas_componentes_m1,
                "horas_componentes_m2": horas_componentes_m2,
                "params": {
                    elem: {
                        "motor": motor_origen_param[elem],
                        "estado": estados_m1[elem] if motor_origen_param[elem] == 1 else estados_m2[elem]
                    } 
                    for elem in elementos 
                    if estados_m1[elem] != "NORMAL" or estados_m2[elem] != "NORMAL"
                }
            })

    # Sección de gráficos
    motor_info = {
        "2CAM3080": {"MT1": "W09060763", "MT2": "W09040922"},
        "2CAM3082": {"MT1": "W09070380", "MT2": ""},
        "2CAM3083": {"MT1": "W08040386", "MT2": "W11030045"},
        "2CAM3085": {"MT1": "W11070959", "MT2": "WX15030035"},
        "2CAM3086": {"MT1": "W11050868", "MT2": "W11090094"},
        "2CAM3087": {"MT1": "W09040921", "MT2": ""},
        "2CAM3090": {"MT1": "W06010096", "MT2": "W13010323"},
        "2CAM3091": {"MT1": "W11060945", "MT2": "W08070121"},
        "2CAM3092": {"MT1": "W12070264", "MT2": "Sin placa"}
    }

    plt.rcParams.update({
        'font.size': 8,
        'font.family': 'Calibri',
        'axes.titlesize': 9,
        'axes.labelsize': 8
    })
    recomendaciones_por_motor = {}  # Diccionario para almacenar recomendaciones por motor

  
    # Dentro de la función generar_informe, en el loop for camion_data in critical_camiones:
# Reemplazar la sección de gráficos con este código:

    for camion_data in critical_camiones:
         numero = camion_data["numero"]
         params_data = camion_data["params"]
    
         # Obtener registros históricos separados
         records_m1 = list(datos_collection.find({
             "camion.numero": numero, 
             "motor_1": True
         }).sort([("fecha_analisis", -1), ("numero_muestra", -1)]))
         
         records_m2 = list(datos_collection.find({
             "camion.numero": numero, 
             "motor_2": True
         }).sort([("fecha_analisis", -1), ("numero_muestra", -1)]))
     
         # Verificar si el camión tiene parámetros que generen gráficos
         tiene_graficos = any(
             param.lower() in ["cfe", "fe", "si", "cu", "cr", "ni"] and param_info["estado"] in ["MONITOREO", "ACCIÓN REQUERIDA"]
             for param, param_info in params_data.items()
         )
         # Si no existen parámetros críticos para graficar, se omite la generación de la tabla y gráficos
         if not tiene_graficos:
             continue
     
         # Crear página y título del análisis
         doc.add_page_break()
         titulo = doc.add_paragraph()
         titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
         run = titulo.add_run(f"ANÁLISIS DEL CAMIÓN {numero}")
         run.bold = True
         run.font.size = Pt(14)
     
         # Crear tabla de información de motores (solo se crea si hay gráficos a generar)
         tabla_info = doc.add_table(rows=4, cols=2)
         tabla_info.style = 'Table Grid'
     
         # Motor 1
         celda = tabla_info.cell(0, 0)
         celda.text = "Motor 1"
         celda = tabla_info.cell(1, 0)
         celda.text = f"N° Serie: {motor_info.get(numero, {}).get('MT1', 'N/A')}"
         celda = tabla_info.cell(2, 0)
         celda.text = f"Estado: {camion_data['estado_m1']}"
         celda = tabla_info.cell(3, 0)
         horas_componentes_m1 = (camion_data['horas_componentes_m1']
                                  if camion_data['horas_componentes_m1'] not in ['N/A', 'Sin registro', '0', 'no hay registros']
                                  else 'Actualizar horometro')
         celda.text = f"Horómetro: {horas_componentes_m1} hrs"
     
         # Motor 2
         celda = tabla_info.cell(0, 1)
         celda.text = "Motor 2"
         celda = tabla_info.cell(1, 1)
         celda.text = f"N° Serie: {motor_info.get(numero, {}).get('MT2', 'N/A')}"
         celda = tabla_info.cell(2, 1)
         celda.text = f"Estado: {camion_data['estado_m2']}"
         celda = tabla_info.cell(3, 1)
         horas_componentes_m2 = (camion_data['horas_componentes_m2']
                                  if camion_data['horas_componentes_m2'] not in ['N/A', 'Sin registro', '0', 'no hay registros']
                                  else 'Actualizar horometro')
         celda.text = f"Horómetro: {horas_componentes_m2} hrs"
     
         # Resaltar motor afectado (se selecciona el primero que cumpla la condición en los parámetros)
         # Reemplazar el bloque donde se define "motor_afectado" con:

         motores_con_parametros_criticos = set()
         
         # Iterar sobre todos los parámetros para detectar motores críticos
         for param, param_info in params_data.items():
             if (
                 param.lower() in ["cfe", "fe", "si", "cu", "cr", "ni"]
                 and param_info["estado"] in ["MONITOREO", "ACCIÓN REQUERIDA"]
             ):
                 motores_con_parametros_criticos.add(param_info["motor"])
         
         # Resaltar TODOS los motores críticos en la tabla
         for motor in motores_con_parametros_criticos:
             motor_celda = tabla_info.cell(0, motor - 1)
             estado_celda = tabla_info.cell(2, motor - 1)
             color = "FFFF00" if camion_data[f"estado_m{motor}"] == "MONITOREO" else "FF0000"
             set_cell_shading(motor_celda, color)
             set_cell_shading(estado_celda, color)
              
         for param, param_info in params_data.items():
            param_lower = param.lower()
            if param_lower in ["cfe", "fe", "si", "cu", "cr", "ni"] and param_info["estado"] in ["MONITOREO", "ACCIÓN REQUERIDA"]:
                motor_origen = param_info["motor"]
                
                # Seleccionar registros del motor específico
                registros_motor = records_m1 if motor_origen == 1 else records_m2
                
                # Verificar si hay datos para este motor
                if not registros_motor:
                    print(f"⚠️ No hay registros para Motor {motor_origen} (Parámetro: {param})")
                    continue
                
                # Procesar registros únicos
                registros_unicos = {}
                for rec in registros_motor:
                    fecha = rec.get("fecha_analisis", "")[:10]
                    param_value = rec.get(param_lower, None)
                    
                    if param_value is None:
                        print(f"⚠️ Parámetro '{param_lower}' no existe en registro del {rec['fecha_analisis']}")
                        continue
                        
                    if fecha not in registros_unicos or rec.get("numero_muestra", 0) > registros_unicos[fecha].get("numero_muestra", 0):
                        registros_unicos[fecha] = rec
                
                # Verificar datos procesados
                if not registros_unicos:
                    print(f"⚠️ No hay datos válidos para {param} en Motor {motor_origen}")
                    continue
              
                registros_ordenados = sorted(
                    registros_unicos.values(),
                    key=lambda r: datetime.strptime(r.get("fecha_analisis", "01-01-2000")[:10], "%d-%m-%Y")
                )
     
                 # Tomar últimos 5 registros (ahora estarán ordenados correctamente)
                ultimos_5_registros = registros_ordenados[-5:] if len(registros_ordenados) >= 5 else registros_ordenados
     
                 # Extraer fechas y valores
                fechas = [rec.get("fecha_analisis", "")[:10] for rec in ultimos_5_registros]
                valores = [to_float(rec.get(param.lower(), 0)) for rec in ultimos_5_registros]
                
                 # Generar gráfico para el parámetro (FE o SI)
                fig, ax = plt.subplots(figsize=(6, 3))
                ax.plot(fechas, valores, marker='o', linewidth=1, label=param.upper())
                
                # Líneas de umbral, si existen
                accion = thresholds.get(param.lower(), {}).get("accion")
                monitoreo = thresholds.get(param.lower(), {}).get("monitoreo")
                if accion is not None:
                    ax.axhline(y=accion, color='red', linestyle='--', label="Acción Requerida")
                if monitoreo is not None:
                    ax.axhline(y=monitoreo, color='yellow', linestyle='--', label="Monitoreo")
                
                ax.set_title(f"Tendencia de {param.upper()} (Motor {motor_origen})")
                ax.legend(loc='upper left', prop={'size': 8})
                ax.grid(True, linestyle=':')
                plt.xticks(rotation=45)
                plt.tight_layout()
    
                img_stream = BytesIO()
                plt.savefig(img_stream, format='png', dpi=150)
                plt.close()
                doc.add_picture(img_stream, width=Inches(5))
    
                # Agregar recomendaciones específicas para el parámetro
                if param.lower() == "si":
                    doc.add_paragraph(
                        "En caso de continuar la condición o incrementar los niveles de las partículas silicio, se recomienda realizar las siguientes acciones:",
                        style="Heading3"
                    )
                    recomendaciones_silicio = [
                        "- Realizar micro filtrado y/o cambio de aceite según plan de mantenimiento.",
                        "- Realizar chequeo de tapa de llenado de aceite de cárter.",
                        "- Realizar chequeo de tapa de inspección planetarios.",
                        "- Realizar chequeo de mangueras y abrazaderas de respiradero.",
                        "- Revisar el estado y la fecha del último cambio del filtro de respiradero del MT.",
                        "- Realizar chequeo y registro fotográfico del tapón magnético del Carter.",
                        "- Realizar chequeo y registro fotográfico del piñón solar y los dientes de los planetarios.",
                        "- Se recomienda realizar metrología y/o cambio de piñón solar por horas de operación.",
                        "- Según inspección realizar cambio del piñón solar por horas de operación.",
                        "- Se recomienda programar una medición de Backlash y EndPlay."
                    ]
                    for rec in recomendaciones_silicio:
                        doc.add_paragraph(rec, style="List Bullet")
                elif param.lower() == "fe":
                    doc.add_paragraph(
                        "En caso de continuar la condición o incrementar los niveles de las partículas fierro o ferromagnéticas, se recomienda realizar las siguientes acciones:",
                        style="Heading3"
                    )
                    recomendaciones_fierro = [
                        "- Realizar micro filtrado y/o cambio de aceite según plan de mantenimiento.",
                        "- Realizar chequeo y registro fotográfico del tapón magnético del Carter.",
                        "- Realizar chequeo y registro fotográfico del piñón solar y los dientes de los planetarios.",
                    ]
                    motor_origen = param_info.get("motor", 1)
                    motor = "Motor 1" if motor_origen == 1 else "Motor 2"
                    # Obtener las horas de operación y el estado según el motor
                    horas_operacion = camion_data.get("horas_componentes_m1", None) if motor == "Motor 1" else camion_data.get("horas_componentes_m2", None)
                    estado_fe = camion_data.get("estado_m1", "Normal") if motor == "Motor 1" else camion_data.get("estado_m2", "Normal")
    
                    if horas_operacion is not None and isinstance(horas_operacion, (int, float)):
                        if horas_operacion >= 8000:
                            recomendaciones_fierro.extend([
                                "- Se recomienda realizar metrología y/o cambio de piñón solar por horas de operación.",
                                "- Según inspección realizar cambio del piñón solar por horas de operación.",
                                "- Se recomienda programar una medición de Backlash y EndPlay."
                            ])
                        elif estado_fe in ["Monitoreo", "Acción Requerida"]:
                            recomendaciones_fierro.extend([
                                "- Según inspección realizar cambio del piñón solar por horas de operación.",
                                "- Se recomienda programar una medición de Backlash y EndPlay."
                            ])
                    
                    for rec in recomendaciones_fierro:
                        doc.add_paragraph(rec, style="List Bullet")
                elif param.lower() == "cfe":
                    doc.add_paragraph(
                        "En caso de continuar la condición o incrementar los niveles de las partículas fierro o ferromagnéticas, se recomienda realizar las siguientes acciones:",
                        style="Heading3"
                    )
                    recomendaciones_cfe= [
                        "- Realizar micro filtrado y/o cambio de aceite según plan de mantenimiento.",
                        "- Realizar chequeo y registro fotográfico del tapón magnético del Carter.",
                        "- Realizar chequeo y registro fotográfico del piñón solar y los dientes de los planetarios.",
                        "- Se recomienda programar una medición de Backlash y EndPlay. "
                    ]
                    for rec in recomendaciones_cfe:
                        doc.add_paragraph(rec, style="List Bullet")
    
            
     
     
         for row in table_datos.rows[1:]:
             for cell in row.cells:
                 for paragraph in cell.paragraphs:
                     for run in paragraph.runs:
                         run.font.name = "Calibri"
                         run.font.size = Pt(8)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="informe_analisis_aceite.docx"'
    doc.save(response)
    return response